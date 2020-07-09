#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Jun  4 08:38:42 2020

@author: dav_
"""


class RussianExerciserCreator:
    
    # take as input the path to the file and the file name
    def __init__(self, file_path):
        self.p = file_path
    
    # Method 1: preposition substitution from a list
    def preposition_substitution(self, pr = ["по", "из", "у", "для", "на", "в", "от", "с", "к", "со", "ко", "про", "до"]):

        # import file word from location
        import docx2txt
        my_text = docx2txt.process(self.p)
        string_my_text = str(my_text)
        phrases = string_my_text.split('\n')
        
        # delete unnecessary empty sentences, which do not contain any information
        phrases_cleaned = []
        for i in range(len(phrases)):
            if phrases[i] == '' or phrases[i] == ' ':
                pass
            else:
                phrases_cleaned.append(phrases[i])
        
        # sobstitute the prepositions of a list by _____
        prep = list(pr)
        phrases_ready = []
        for i in range(len(phrases_cleaned)):
            words = phrases_cleaned[i].split(' ')
            for j in range(len(words)):
                if words[j] in prep:
                    words[j] = '____'
                else:
                    pass
            for j in range(len(words)):
                if j == 0:
                    new_phrase = words[j]
                else:
                    new_phrase = new_phrase + ' ' + words[j]
            phrases_ready.append(new_phrase)
        
        # create the final string with the ready exercise
        final = 'Complete the text with the following prepositions: '+ str(prep)[1:-1]
        for i in range(len(phrases_ready)):
            if i == 0:
                final = final + '\n\n\n' + phrases_ready[i]
            else:
                final = final + '\n\n' + phrases_ready[i]
        
        from docx import Document
        from docx.shared import Inches
        
        # save file
        document = Document()
        document.add_paragraph(final)
        document.save(self.p[:-4] + '_prepositions.docx')


    # Method 2: words substitution from a list
    def words_substitution(self, ww = []):

        # import file word from location
        import docx2txt
        my_text = docx2txt.process(self.p)
        string_my_text = str(my_text)
        phrases = string_my_text.split('\n')
        
        # delete unnecessary empty sentences, which do not contain any information
        phrases_cleaned = []
        for i in range(len(phrases)):
            if phrases[i] == '' or phrases[i] == ' ':
                pass
            else:
                phrases_cleaned.append(phrases[i])
        
        # substitute the prepositions of a list by _____
        import numpy as np
        import pandas as pd
        words_substitute0 = pd.read_excel(ww)
        words_substitute = list(words_substitute0.iloc[:,0])
        phrases_ready = []
        for i in range(len(phrases_cleaned)):
            words = phrases_cleaned[i].split(' ')
            for j in range(len(words)):
                if words[j][-1] in [',', '.', '!', '?']:
                    if words[j][:-1] in words_substitute:
                        words[j] = '____ ' + words[j][-1]
                    else:
                        pass
                else:
                    if words[j] in words_substitute:
                        words[j] = '____'
                    else:
                        pass
                    
            for j in range(len(words)):
                if j == 0:
                    new_phrase = words[j]
                else:
                    new_phrase = new_phrase + ' ' + words[j]
            phrases_ready.append(new_phrase)
        
        # create the final string with the ready exercise
        final = 'Complete the text with the following words: '+ str(words_substitute)[1:-1]
        for i in range(len(phrases_ready)):
            if i == 0:
                final = final + '\n\n\n' + phrases_ready[i]
            else:
                final = final + '\n\n' + phrases_ready[i]
        
        from docx import Document
        from docx.shared import Inches
        
        # save file
        document = Document()
        document.add_paragraph(final)
        document.save(self.p[:-4] + '_words.docx')

   
    # Method 3: words shuffle exercise
    def words_shuffle(self):

        # import file word from location
        import docx2txt
        my_text = docx2txt.process(self.p)
        string_my_text = str(my_text)
        phrases = string_my_text.split('\n')
        
        # delete unnecessary empty sentences, which do not contain any information
        phrases_cleaned = []
        for i in range(len(phrases)):
            if phrases[i] == '' or phrases[i] == ' ':
                pass
            else:
                phrases_cleaned.append(phrases[i])
                
        # shuffle the words within a sentence
        import random
        phrases_ready = []
        for i in range(len(phrases_cleaned)):
            sentences = phrases_cleaned[i].split('. ')
            sentences[-1] = sentences[-1][:-1]
            for k in range(len(sentences)):
                words = sentences[k].split(' ')
                words_shuffled = random.sample(words, len(words))
                sentence_ready = words_shuffled[0].lower()
                if len(words_shuffled)>0:
                    for j in range(1, len(words_shuffled)):
                        sentence_ready = sentence_ready + ' / ' + words_shuffled[j].lower()
                    phrases_ready.append(sentence_ready)
                else:
                    phrases_ready.append(sentence_ready)
            
        # exsporting file        
        final = 'Order the word in order to obtain reasonable sentences'
        for i in range(len(phrases_ready)):
            if i == 0:
                final = final + '\n\n\n' +  str(i+1) + '. ' + phrases_ready[i]
            else:
                final = final + '\n\n' +  str(i+1) + '. ' + phrases_ready[i]
            
        from docx import Document
        from docx.shared import Inches
        
        # save file
        document = Document()
        document.add_paragraph(final)
        document.save(self.p[:-4] + '_words-shuffle.docx')     
        
            
    # Method 4: verbs tenses
    def verbs_tenses(self, verb_data):
        
        # import file word from location
        import docx2txt
        my_text = docx2txt.process(self.p)
        string_my_text = str(my_text)
        phrases = string_my_text.split('\n')
        
        # delete unnecessary empty sentences, which do not contain any information
        phrases_cleaned = []
        for i in range(len(phrases)):
            if phrases[i] == '' or phrases[i] == ' ':
                pass
            else:
                phrases_cleaned.append(phrases[i])
        
        # list of lists of phrases words
        phrases_listed = []
        for i in range(len(phrases_cleaned)):
            phrases_listed.append(phrases_cleaned[i].split(' '))
         
        words_ready_list = []
        fu = ['буду', 'будешь', 'будет', 'будем', 'будете', 'будут']
        
        verbs_list = []
        for i in range(len(verb_data)):
            verb_cog = []
            for j in range(len(verb_data.columns)-1):
                verb_cog.append(verb_data.iloc[i,j])
            verbs_list.append(verb_cog)
            
        for i in range(len(phrases_listed)):
            for r in range(len(verbs_list)):
                if phrases_listed[i][0].lower() in verbs_list[r]:
                    aaa = ['________ (' + verbs_list[r][0] +')']
                    break
                if r == len(verbs_list) - 1:
                            aaa = [(phrases_listed[i][0])]
                        
            for j in range(1, len(phrases_listed[i])):
                for k in range(len(verbs_list)):
                    # if there is , . ! ? 
                    if len(phrases_listed[i][j]) > 0:
                        if phrases_listed[i][j][-1] in [',', '.', '!', '?']:
                            if phrases_listed[i][j][:-1].lower() in verbs_list[k]:
                                if phrases_listed[i][j-1].lower() in fu:
                                    aaa[-1] = '________ (' + verbs_list[k][0] +')' + phrases_listed[i][j][-1]
                                    break
                                else:
                                    aaa.append('________ (' + verbs_list[k][0] +')' + phrases_listed[i][j][-1])
                                    break
                            
                                
                            if k == len(verbs_list) - 1:
                                aaa.append(phrases_listed[i][j])
    
                        else:
                            if phrases_listed[i][j].lower() in verbs_list[k]:
                                if phrases_listed[i][j-1].lower() in fu:
                                    aaa[-1] = '________ (' + verbs_list[k][0] +')'
                                    break
                                else:
                                    aaa.append('________ (' + verbs_list[k][0] +')')
                                    break
                 
                            if k == len(verbs_list) - 1:
                                aaa.append(phrases_listed[i][j])
            words_ready_list.append(aaa)
                    
        phrase_ready = []
        for i in range(len(words_ready_list)):
            phrase_ready.append(' '.join(words_ready_list[i]))
        
        
        final = 'Put verbs between brackets in the correct form.'
        for i in range(len(phrase_ready)):
            if i == 0:
                final = final + '\n\n\n' + phrase_ready[i]
            else:
                final = final + '\n\n' + phrase_ready[i]
        
        from docx import Document
        from docx.shared import Inches
        
        # save file
        document = Document()
        document.add_paragraph(final)
        document.save(self.p[:-4] + '_verbs-tenses.docx')


    # Method 5: imperfective/perfective verbs ALL
    def sv_nsv_all(self, verb_data, verbs_pair):
        
        # import file word from location
        import docx2txt
        my_text = docx2txt.process(self.p)
        string_my_text = str(my_text)
        phrases = string_my_text.split('\n')
        
        # delete unnecessary empty sentences, which do not contain any information
        phrases_cleaned = []
        for i in range(len(phrases)):
            if phrases[i] == '' or phrases[i] == ' ':
                pass
            else:
                phrases_cleaned.append(phrases[i])
        
        # list of lists of phrases words
        phrases_listed = []
        for i in range(len(phrases_cleaned)):
            phrases_listed.append(phrases_cleaned[i].split(' '))
         
        words_ready_list = []
        fu = ['буду', 'будешь', 'будет', 'будем', 'будете', 'будут']
        
        verbs_list = []
        for i in range(len(verb_data)):
            verb_cog = []
            for j in range(len(verb_data.columns)-1):
                verb_cog.append(verb_data.iloc[i,j])
            verbs_list.append(verb_cog)
            
        sns = verbs_pair
        for i in range(len(phrases_listed)):
            for r in range(len(verbs_list)):
                if phrases_listed[i][0].lower() in verbs_list[r]:
                    for rr in range(len(sns)):
                        if verbs_list[r][0].lower() == sns.iloc[rr, 0] or verbs_list[r][0].lower() == sns.iloc[rr, 1]:
                            aaa = ['________ (' + sns.iloc[rr, 0] +' / ' + sns.iloc[rr, 1] +')']
                            break
                        elif rr == len(sns) -1:
                            aaa = ['________ (' + verbs_list[r][0] +')']
                            break
                if r == len(verbs_list) - 1:
                    aaa = [(phrases_listed[i][0])]
                        
            for j in range(1, len(phrases_listed[i])):
                for k in range(len(verbs_list)):
                    # if there is , . ! ? 
                    if len(phrases_listed[i][j]) > 0:
                        if phrases_listed[i][j][-1] in [',', '.', '!', '?']:
                            if phrases_listed[i][j][:-1].lower() in verbs_list[k]:
                                for rr in range(2, len(sns)):
                                    if verbs_list[k][0].lower() == sns.iloc[rr, 0] or verbs_list[k][0].lower() == sns.iloc[rr, 1]:
                                        if phrases_listed[i][j-1].lower() in fu:
                                            aaa[-1] = '________ (' + sns.iloc[rr, 0] + ' / ' + sns.iloc[rr, 1] +')' + phrases_listed[i][j][-1]
                                            break
                                        else:
                                            aaa.append('________ (' + sns.iloc[rr, 0] + ' / ' + sns.iloc[rr, 1] +')' + phrases_listed[i][j][-1])
                                            break
                                        
                                    elif rr == len(sns) -1:
                                        if phrases_listed[i][j-1].lower() in fu:
                                            aaa[-1] = '________ (' + verbs_list[k][0] +')' + phrases_listed[i][j][-1]
                                            break
                                        else:
                                            aaa.append('________ (' + verbs_list[k][0] +')' + phrases_listed[i][j][-1])
                                            break
                                    
                                break
                            
                            if k == len(verbs_list) - 1:
                                aaa.append(phrases_listed[i][j])
    
                        else:
                            if phrases_listed[i][j].lower() in verbs_list[k]:
                                for rr in range(2, len(sns)):
                                    if verbs_list[k][0].lower() == sns.iloc[rr, 0] or verbs_list[k][0].lower() == sns.iloc[rr, 1]:
                                        if phrases_listed[i][j-1].lower() in fu:
                                            aaa[-1] = '________ (' + sns.iloc[rr, 0] + ' / ' + sns.iloc[rr, 1] +')'
                                            break
                                        
                                        else:
                                            aaa.append('________ (' + sns.iloc[rr, 0] + ' / ' + sns.iloc[rr, 1] +')')
                                            break
                                        
                                    elif rr == len(sns) -1:
                                        if phrases_listed[i][j-1].lower() in fu:
                                            aaa[-1] = '________ (' + verbs_list[k][0] +')'
                                            break
                                        else:
                                            aaa.append('________ (' + verbs_list[k][0] +')')
                                            break

                                break
                 
                            if k == len(verbs_list) - 1:
                                aaa.append(phrases_listed[i][j])
            words_ready_list.append(aaa)
                    
        phrases_ready = []
        for i in range(len(words_ready_list)):
            phrases_ready.append(' '.join(words_ready_list[i]))
        
        
        final = 'Choose the right verb form and put it in the correct tense'
        for i in range(len(phrases_ready)):
            if i == 0:
                final = final + '\n\n\n' + phrases_ready[i]
            else:
                final = final + '\n\n' + phrases_ready[i]
        
        from docx import Document
        from docx.shared import Inches
       
        # save file
        document = Document()
        document.add_paragraph(final)
        document.save(self.p[:-4] + '_sv-nsv-all.docx')

        
    # Method 6: imperfective/perfective verbs ONLY
    def sv_nsv_only(self, verb_data, verbs_pair):
        
        # import file word from location
        import docx2txt
        my_text = docx2txt.process(self.p)
        string_my_text = str(my_text)
        phrases = string_my_text.split('\n')
        
        # delete unnecessary empty sentences, which do not contain any information
        phrases_cleaned = []
        for i in range(len(phrases)):
            if phrases[i] == '' or phrases[i] == ' ':
                pass
            else:
                phrases_cleaned.append(phrases[i])
        
        # list of lists of phrases words
        phrases_listed = []
        for i in range(len(phrases_cleaned)):
            phrases_listed.append(phrases_cleaned[i].split(' '))
         
        words_ready_list = []
        fu = ['буду', 'будешь', 'будет', 'будем', 'будете', 'будут']
        
        verbs_list = []
        for i in range(len(verb_data)):
            verb_cog = []
            for j in range(len(verb_data.columns)-1):
                verb_cog.append(verb_data.iloc[i,j])
            verbs_list.append(verb_cog)
            
        sns = verbs_pair
        for i in range(len(phrases_listed)):
            for r in range(len(verbs_list)):
                if phrases_listed[i][0].lower() in verbs_list[r]:
                    for rr in range(2, len(sns)):
                        if verbs_list[r][0].lower() == sns.iloc[rr, 0] or verbs_list[r][0].lower() == sns.iloc[rr, 1]:
                            aaa = ['________ (' + sns.iloc[rr, 0] +' / ' + sns.iloc[rr, 1] +')']
                            break
                        elif rr == len(sns) -1:
                            aaa = ['________ (' + verbs_list[r][0] +')']
                            break
                if r == len(verbs_list) - 1:
                    aaa = [(phrases_listed[i][0])]
                        
            for j in range(1, len(phrases_listed[i])):
                for k in range(len(verbs_list)):
                    # if there is , . ! ? 
                    if len(phrases_listed[i][j]) > 0:
                        if phrases_listed[i][j][-1] in [',', '.', '!', '?']:
                            if phrases_listed[i][j][:-1].lower() in verbs_list[k]:
                                for rr in range(2, len(sns)):
                                    if verbs_list[k][0].lower() == sns.iloc[rr, 0] or verbs_list[k][0].lower() == sns.iloc[rr, 1]:
                                        if phrases_listed[i][j-1].lower() in fu:
                                            aaa[-1] = '________ (' + sns.iloc[rr, 0] + ' / ' + sns.iloc[rr, 1] +')' + phrases_listed[i][j][-1]
                                            break
                                        else:
                                            aaa.append('________ (' + sns.iloc[rr, 0] + ' / ' + sns.iloc[rr, 1] +')' + phrases_listed[i][j][-1])
                                            break
                                        
                                    elif rr == len(sns) -1:
                                        aaa.append(phrases_listed[i][j])
                                    
                                break
                            
                            if k == len(verbs_list) - 1:
                                aaa.append(phrases_listed[i][j])
    
                        else:
                            if phrases_listed[i][j].lower() in verbs_list[k]:
                                for rr in range(2, len(sns)):
                                    if verbs_list[k][0].lower() == sns.iloc[rr, 0] or verbs_list[k][0].lower() == sns.iloc[rr, 1]:
                                        if phrases_listed[i][j-1].lower() in fu:
                                            aaa[-1] = '________ (' + sns.iloc[rr, 0] + ' / ' + sns.iloc[rr, 1] +')'
                                            break
                                        
                                        else:
                                            aaa.append('________ (' + sns.iloc[rr, 0] + ' / ' + sns.iloc[rr, 1] +')')
                                            break
                                        
                                    elif rr == len(sns) -1:
                                        aaa.append(phrases_listed[i][j])

                                break
                 
                            if k == len(verbs_list) - 1:
                                aaa.append(phrases_listed[i][j])
            words_ready_list.append(aaa)
                    
        phrases_ready = []
        for i in range(len(words_ready_list)):
            phrases_ready.append(' '.join(words_ready_list[i]))
        
        
        final = 'Decline the verbs'
        for i in range(len(phrases_ready)):
            if i == 0:
                final = final + '\n\n\n'  + phrases_ready[i]
            else:
                final = final + '\n\n' + phrases_ready[i]
        
        from docx import Document
        from docx.shared import Inches
       
        # save file
        document = Document()
        document.add_paragraph(final)
        document.save(self.p[:-4] + '_sv-nsv-only.docx')
        

# import needed libraries
import tkinter as tk
from tkinter import filedialog
import numpy as np
import pandas as pd

# create a prompt for choosing a word file
print('choose the word (.docx) file, which you want to convert into an exercise.')
root = tk.Tk()
root.withdraw()
file_path = filedialog.askopenfilename()

# ask the user to select an exercise
methods = input('Which exercise do you want to create?' + '\n' +'Write:' + '\n' +
                '1 for preposition substitution'+ '\n' + '2 for words substitution'+ '\n' + 
                '3 for words shuffle' + '\n' + '4 for verbs tenses' + '\n' + '5 for nsv-sv all' + '\n' + '6 for nsv-sv only'
                + '\n\n your choice:  ')

# perform the related mathod selected from the user
if methods == '1':
    exercise = RussianExerciserCreator(file_path)
    exercise.preposition_substitution()

if methods == '2':
    print('choose the words file to substitute')
    root1 = tk.Tk()
    root1.withdraw()
    ww = filedialog.askopenfilename()
    exercise = RussianExerciserCreator(file_path)
    exercise.words_substitution(ww)

if methods == '3':
    exercise = RussianExerciserCreator(file_path)
    exercise.words_shuffle()

if methods == '4':
    print('choose the verb file')
    root4 = tk.Tk()
    root4.withdraw()
    verbs = pd.read_excel(filedialog.askopenfilename())
    exercise = RussianExerciserCreator(file_path)
    exercise.verbs_tenses(verbs)

if methods == '5':
    print('choose the verb file')
    root5 = tk.Tk()
    root5.withdraw()
    verbs = pd.read_excel(filedialog.askopenfilename())
    print('choose the verb pair file')
    root5 = tk.Tk()
    root5.withdraw()
    verbs_pairs = pd.read_excel(filedialog.askopenfilename())
    exercise = RussianExerciserCreator(file_path)
    exercise.sv_nsv_all(verbs, verbs_pairs)

if methods == '6':
    print('choose the verb file')
    root6 = tk.Tk()
    root6.withdraw()
    verbs = pd.read_excel(filedialog.askopenfilename())
    print('choose the verb pair file')
    root6 = tk.Tk()
    root6.withdraw()
    verbs_pairs = pd.read_excel(filedialog.askopenfilename())
    exercise = RussianExerciserCreator(file_path)
    exercise.sv_nsv_only(verbs, verbs_pairs)





                
